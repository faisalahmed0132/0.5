
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.keys import Keys
import time
from docx import Document
from docx.shared import Pt

# إعداد المتصفح
options = webdriver.ChromeOptions()
options.add_argument("--start-maximized")

driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
driver.get("https://teams.microsoft.com")

input("➡️ بعد تسجيل الدخول والوصول للقناة المطلوبة يدويًا، اضغط Enter للمتابعة...")

# تمرير تلقائي للأعلى للحصول على منشورات قديمة
scroll_pause = 2
scrolls = 25  # عدل هذا الرقم إذا أردت الرجوع لأكثر من شهرين
chat_container = driver.find_element(By.TAG_NAME, 'body')

for _ in range(scrolls):
    chat_container.send_keys(Keys.HOME)
    time.sleep(scroll_pause)

# استخراج المنشورات
posts = driver.find_elements(By.CLASS_NAME, "message-body-content")

# تجهيز ملف Word
doc = Document()
doc.add_heading('منشورات قناة Microsoft Teams – آخر شهرين', 0)

for i, post in enumerate(posts, 1):
    text = post.text.strip()
    if text:
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {text}")
        run.font.size = Pt(11)

doc.save("teams_posts_last_2_months.docx")
print("✅ تم حفظ الملف: teams_posts_last_2_months.docx")

driver.quit()
